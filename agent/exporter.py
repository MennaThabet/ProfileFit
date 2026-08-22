"""
Deterministic CV Exporter (Exporter Agent).

Takes a structured TailoredCV (from selector_tailor.generate_tailored_cv)
plus the candidate's CONTACT ProfileItem(s) (from
profile_parser.get_contact_items) and produces a consistently-styled .docx,
then optionally converts it to .pdf.

This is intentionally NOT an LLM call.
"""

import shutil
import subprocess
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from schemas import ItemType, ProfileItem, TailoredCV, TailoredSectionItem

# Fixed document style constants 
FONT_NAME = "Calibri"
NAME_SIZE = Pt(20)
HEADING_SIZE = Pt(13)
BODY_SIZE = Pt(10.5)
CONTACT_SIZE = Pt(10)
ACCENT_COLOR = RGBColor(0x1F, 0x3B, 0x57)  
BODY_COLOR = RGBColor(0x20, 0x20, 0x20)
WARNING_COLOR = RGBColor(0xB0, 0x2A, 0x2A)

# Section rendering order and display headings.
SECTION_ORDER: list[tuple[ItemType, str]] = [
    (ItemType.EDUCATION, "Education"),
    (ItemType.EXPERIENCE, "Experience"),
    (ItemType.PROJECT, "Projects"),
    (ItemType.CERTIFICATION, "Certifications"),
    (ItemType.SKILL, "Skills"),
    (ItemType.ACHIEVEMENT, "Achievements"),
]


def _set_base_styles(document: Document) -> None:
    """Fix the Normal style so every paragraph defaults to the same font/size."""
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BODY_COLOR


def _add_heading(document: Document, text: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text.upper())
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR
    pPr = heading._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3B57")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _extract_contact_fields(contact_items: list[ProfileItem]) -> dict[str, str]:
    """
    Best-effort split of a CONTACT ProfileItem's free-text description into
    a name + a line of contact details. 
    """
    if not contact_items:
        return {"name": "", "details": ""}
    primary = contact_items[0]
    return {"name": primary.title, "details": primary.description}


def _add_contact_header(document: Document, contact_items: list[ProfileItem]) -> None:
    fields = _extract_contact_fields(contact_items)

    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(fields["name"] or "Candidate Name")
    name_run.font.name = FONT_NAME
    name_run.font.size = NAME_SIZE
    name_run.font.bold = True
    name_run.font.color.rgb = ACCENT_COLOR

    if fields["details"]:
        contact_para = document.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(fields["details"])
        contact_run.font.name = FONT_NAME
        contact_run.font.size = CONTACT_SIZE
        contact_run.font.color.rgb = BODY_COLOR
    elif not contact_items:
        warn_para = document.add_paragraph()
        warn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warn_run = warn_para.add_run(
            "[No contact information available — pass contact_items from the "
            "parsed master profile]"
        )
        warn_run.font.name = FONT_NAME
        warn_run.font.size = CONTACT_SIZE
        warn_run.font.italic = True
        warn_run.font.color.rgb = WARNING_COLOR


def _add_summary(document: Document, summary: str) -> None:
    if not summary:
        return
    _add_heading(document, "Professional Summary")
    para = document.add_paragraph()
    run = para.add_run(summary)
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE
    run.font.color.rgb = BODY_COLOR


def _add_section_item(document: Document, item: TailoredSectionItem) -> None:
    title_para = document.add_paragraph()
    title_para.paragraph_format.space_before = Pt(6)
    title_text = item.title
    if item.dates:
        title_text = f"{item.title}  |  {item.dates}"
    title_run = title_para.add_run(title_text)
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    title_run.font.color.rgb = BODY_COLOR

    for bullet in item.tailored_bullets:
        bullet_para = document.add_paragraph(style="List Bullet")
        bullet_para.paragraph_format.space_after = Pt(2)
        bullet_run = bullet_para.add_run(bullet.text)
        bullet_run.font.name = FONT_NAME
        bullet_run.font.size = BODY_SIZE
        bullet_run.font.color.rgb = BODY_COLOR


def _add_empty_sections_notice(document: Document) -> None:
    _add_heading(document, "Notice")
    para = document.add_paragraph()
    run = para.add_run(
        "No relevant profile items were found for this job posting. This CV "
        "could not be tailored with the available master profile data. "
        "Review the missing_skills field and consider adding more detail to "
        "the master profile."
    )
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE
    run.font.italic = True
    run.font.color.rgb = WARNING_COLOR


def render_cv_docx(
    tailored_cv: TailoredCV,
    contact_items: Optional[list[ProfileItem]] = None,
    output_path: str | Path = "output/tailored_cv.docx",
) -> Path:
    document = Document()
    _set_base_styles(document)

    _add_contact_header(document, contact_items or [])
    _add_summary(document, tailored_cv.professional_summary)

    by_section: dict[ItemType, list[TailoredSectionItem]] = {}
    for item in tailored_cv.sections:
        by_section.setdefault(item.section, []).append(item)

    if not tailored_cv.sections:
        _add_empty_sections_notice(document)
    else:
        for item_type, heading in SECTION_ORDER:
            items = by_section.get(item_type)
            if not items:
                continue
            _add_heading(document, heading)
            for item in sorted(items, key=lambda x: x.relevance_score, reverse=True):
                _add_section_item(document, item)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Rendered CV (.docx) saved to {output_path}")
    return output_path


def _find_soffice() -> Optional[str]:
    for name in ("soffice", "soffice.exe", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    default_win = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if default_win.exists():
        return str(default_win)
    return None


def convert_docx_to_pdf(docx_path: str | Path, output_dir: Optional[str | Path] = None) -> Optional[Path]:
    soffice = _find_soffice()
    if soffice is None:
        print(
            "Warning: LibreOffice not found — skipping PDF conversion. "
            "The .docx was still saved. Install LibreOffice to enable PDF export, "
            "or convert manually (Word/Google Docs > Export as PDF).",
            file=sys.stderr,
        )
        return None

    docx_path = Path(docx_path)
    out_dir = Path(output_dir) if output_dir else docx_path.parent

    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        print(f"Warning: PDF conversion failed: {e}", file=sys.stderr)
        return None

    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if pdf_path.exists():
        print(f"Rendered CV (.pdf) saved to {pdf_path}")
        return pdf_path

    print("Warning: PDF conversion ran but no output file was found.", file=sys.stderr)
    return None


def export_cv(
    tailored_cv: TailoredCV,
    contact_items: Optional[list[ProfileItem]] = None,
    output_path: str | Path = "output/tailored_cv.docx",
    also_pdf: bool = True,
) -> dict[str, Optional[Path]]:
    docx_path = render_cv_docx(tailored_cv, contact_items=contact_items, output_path=output_path)
    pdf_path = convert_docx_to_pdf(docx_path) if also_pdf else None
    return {"docx_path": docx_path, "pdf_path": pdf_path}


if __name__ == "__main__":
    import argparse

    from profile_parser import get_contact_items, parse_profile_data
    from requirement_extractor import extract_requirements
    from selector_tailor import generate_tailored_cv, job_requirements_from_extraction

    parser = argparse.ArgumentParser(
        description="Run the full parse -> extract -> tailor -> export pipeline for one posting."
    )
    parser.add_argument("--text", help="Raw job posting text.")
    parser.add_argument("--file", help="Path to a PDF containing the job posting.")
    parser.add_argument("--url", help="URL of the job posting to fetch.")
    parser.add_argument(
        "--profile-instruction",
        help='Instruction for the profile parser, e.g. "github link is https://github.com/USERNAME". '
        "Pass this (or --profile-file) so the export includes real contact info instead "
        "of the placeholder warning.",
    )
    parser.add_argument("--profile-file", help="Path to a resume/CV PDF to parse instead.")
    parser.add_argument("--output", default="output/tailored_cv.docx", help="Output .docx path.")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF conversion.")
    args = parser.parse_args()

    if args.url:
        posting_source, fetch_first = f"job posting is: {args.url}", True
    elif args.file:
        posting_source, fetch_first = f"pdf path is {args.file}", True
    elif args.text:
        posting_source, fetch_first = args.text, False
    else:
        posting_source, fetch_first = (
            """
            Data Analyst Intern — Acme Corp
            Requirements:
            - Bachelor's degree in progress (CS, Statistics, or related field)
            - Required: proficiency in SQL and Python (pandas)
            - Preferred: experience with Tableau or Power BI
            - Nice to have: exposure to A/B testing concepts
            """,
            False,
        )

    if args.profile_file:
        profile_instruction = (
            f"read the cv file at this path: {args.profile_file}, "
            f"and extract the profile data from it."
        )
    elif args.profile_instruction:
        profile_instruction = args.profile_instruction
    else:
        profile_instruction = None

    try:
        if profile_instruction:
            print("Parsing master profile...")
            profile_items = parse_profile_data(profile_instruction, save=False)
            contact_items = get_contact_items(profile_items)
            print(f"Parsed {len(profile_items)} items, {len(contact_items)} CONTACT item(s).\n")
        else:
            print(
                "Warning: no --profile-instruction or --profile-file given — "
                "falling back to RAG retrieval for candidate context, and the "
                "exported document will show the 'no contact information' "
                "placeholder. Pass --profile-instruction/--profile-file for a "
                "complete run.\n"
            )
            profile_items = None
            contact_items = []

        extraction = extract_requirements(posting_source, fetch_first=fetch_first, save=False)
        if extraction.flagged_for_review:
            print(f"[FLAGGED] {extraction.flag_reason}")
            sys.exit(1)

        job_reqs = job_requirements_from_extraction(extraction)
        print(f"Tailoring for: {job_reqs.job_title}")

        if profile_items is not None:
            tailored_cv = generate_tailored_cv(job_reqs, candidate_items=profile_items)
        else:
            import json
            from rag import rag_search

            rag_result = json.loads(
                rag_search(
                    f"Candidate profile experiences, projects, education, and skills "
                    f"relevant to {job_reqs.job_title}. Required: "
                    f"{', '.join(job_reqs.required_skills)}.",
                    top_k=10,
                    source_filter="master_profiles",
                )
            )
            tailored_cv = generate_tailored_cv(job_reqs, rag_sources=rag_result.get("sources", []))

        export_cv(
            tailored_cv,
            contact_items=contact_items,
            output_path=args.output,
            also_pdf=not args.no_pdf,
        )

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)